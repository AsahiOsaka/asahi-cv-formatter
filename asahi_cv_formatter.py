# Enhanced Asahi CV Formatter - Professional Design
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import fitz  # PyMuPDF
import re
from PIL import Image

# --- Custom CSS for Professional Styling ---
def apply_custom_css():
    st.markdown("""
    <style>
    /* Main app styling */
    .main {
        padding-top: 2rem;
        padding-bottom: 2rem;
    }
    
    /* Header styling */
    .main-header {
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        padding: 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    .main-header h1 {
        color: white;
        text-align: center;
        margin: 0;
        font-weight: 700;
        font-size: 2.5rem;
    }
    
    .main-header p {
        color: #e2e8f0;
        text-align: center;
        margin: 0.5rem 0 0 0;
        font-size: 1.1rem;
    }
    
    /* Upload section styling */
    .upload-section {
        background: #f8fafc;
        padding: 2rem;
        border-radius: 12px;
        border: 2px dashed #cbd5e1;
        margin-bottom: 2rem;
        transition: all 0.3s ease;
    }
    
    .upload-section:hover {
        border-color: #3b82f6;
        background: #f1f5f9;
    }
    
    /* Input styling */
    .stTextInput > div > div > input {
        border: 2px solid #e2e8f0;
        border-radius: 8px;
        padding: 0.75rem;
        font-size: 1rem;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #3b82f6;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1);
    }
    
    .stNumberInput > div > div > input {
        border: 2px solid #e2e8f0;
        border-radius: 8px;
        padding: 0.75rem;
        font-size: 1rem;
    }
    
    /* Button styling */
    .stButton > button {
        background: linear-gradient(90deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        font-size: 1.1rem;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(59, 130, 246, 0.4);
    }
    
    /* Download button styling */
    .stDownloadButton > button {
        background: linear-gradient(90deg, #059669 0%, #10b981 100%);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        font-size: 1.1rem;
        width: 100%;
    }
    
    /* Success message styling */
    .success-message {
        background: linear-gradient(90deg, #059669 0%, #10b981 100%);
        color: white;
        padding: 1rem 2rem;
        border-radius: 8px;
        margin: 1rem 0;
        text-align: center;
        font-weight: 600;
    }
    
    /* Preview section */
    .preview-section {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    /* Info boxes */
    .info-box {
        background: #eff6ff;
        border-left: 4px solid #3b82f6;
        padding: 1rem 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }
    
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """, unsafe_allow_html=True)

# --- Helper: Extract text from PDF ---
def extract_text_from_pdf(file):
    text = ""
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""
    return text

# --- Helper: Extract text from DOCX ---
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

# --- Helper: Abbreviate Name and Add Age ---
def abbreviate_name_age(full_name, age):
    try:
        name_parts = [part.strip() for part in full_name.strip().split() if part.strip()]
        if not name_parts:
            return f"N.A.{age}yrs"
        
        initials = ''.join([part[0].upper() + '.' for part in name_parts])
        return f"{initials} {age}yrs"
    except Exception:
        return f"N.A.{age}yrs"

# --- Helper: Add header with logo to every page ---
def add_header_with_logo(doc, logo_img):
    # Create header section
    section = doc.sections[0]
    header = section.header
    
    # Clear any existing header content
    for paragraph in header.paragraphs:
        paragraph.clear()
    
    # Create paragraph with absolute right alignment
    logo_para = header.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add tab stop at far right margin
    tab_stops = logo_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Add tab character to force right alignment
    logo_run = logo_para.add_run("\t")
    
    # Add logo at exact original size (2.634" x 0.508")
    image_stream = BytesIO()
    logo_img.save(image_stream, format='PNG')
    image_stream.seek(0)
    logo_run.add_picture(image_stream, width=Inches(2.634), height=Inches(0.508))
    
    # Adjust header distance from top
    section.header_distance = Inches(0.4)  # Bring header closer to top edge


# --- Helper: Generate professional CV using enhanced Asahi format ---
def generate_asahi_cv(raw_text, logo_img, candidate_name, age):
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.2)  # Increased top margin for header
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Add logo to header (appears on every page)
    add_header_with_logo(doc, logo_img)
    
    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Add centered name section (in document body, not header)
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(24)  # Add spacing after name
    
    name_run = name_paragraph.add_run(abbreviate_name_age(candidate_name, age))
    name_run.font.name = 'ＭＳ 明朝'  # MS Mincho Body Asian font
    name_run.font.size = Pt(16)  # Professional size
    name_run.font.bold = True
    
    # Add spacing after name
    doc.add_paragraph()
    
    # Process and add content with simple, clean formatting
    content_lines = [line.strip() for line in raw_text.strip().split("\n") if line.strip()]
    
    for line in content_lines:
        if line.strip():
            doc.add_paragraph(line.strip())
    
    return doc

# --- Main Streamlit App ---
def main():
    st.set_page_config(
        page_title="Asahi CV Formatter", 
        layout="centered",
        initial_sidebar_state="collapsed"
    )
    
    # Apply custom styling
    apply_custom_css()
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🏢 Asahi CV Formatter</h1>
        <p>Transform candidate CVs into Asahi format</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Upload section
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        uploaded_file = st.file_uploader(
            "📄 Upload Candidate CV", 
            type=["docx", "pdf"],
            help="Upload a DOCX or PDF file containing the candidate's CV"
        )
    
    with col2:
        st.markdown("**Supported formats:**")
        st.markdown("• PDF documents")
        st.markdown("• DOCX documents")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Input fields
    col1, col2 = st.columns(2)
    
    with col1:
        candidate_name = st.text_input(
            "👤 Candidate Full Name", 
            placeholder="e.g., Ray Dooman",
            help="Enter the candidate's complete name"
        )
    
    with col2:
        age = st.number_input(
            "🎂 Candidate Age", 
            min_value=18, 
            max_value=99, 
            step=1,
            help="Enter the candidate's age (18-99)"
        )
    
    # Process file if all inputs are provided
    if uploaded_file and candidate_name and age:
        try:
            # Load logo
            logo_path = "asahi_logo-04.jpg"
            try:
                logo_img = Image.open(logo_path)
            except FileNotFoundError:
                st.error("⚠️ Logo file 'asahi_logo-04.jpg' not found. Please ensure it's in the same directory.")
                st.stop()
            
            # Extract text based on file type
            with st.spinner("📖 Extracting text from uploaded file..."):
                if uploaded_file.name.lower().endswith(".pdf"):
                    raw_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.name.lower().endswith(".docx"):
                    raw_text = extract_text_from_docx(uploaded_file)
                else:
                    st.error("❌ Unsupported file type. Please upload a PDF or DOCX file.")
                    st.stop()
            
            if not raw_text.strip():
                st.error("❌ No text could be extracted from the file. Please check the file format.")
                st.stop()
            
            # Display preview in professional container
            st.markdown("### 📋 Extracted CV Preview")
            st.markdown('<div class="preview-section">', unsafe_allow_html=True)
            
            preview_text = raw_text[:2000] + "..." if len(raw_text) > 2000 else raw_text
            st.text_area(
                "Content Preview:", 
                preview_text, 
                height=250,
                help="Preview of extracted text from your CV file"
            )
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Show candidate info preview
            st.markdown("### 👤 Candidate Information")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.info(f"**Name:** {candidate_name}")
            with col2:
                st.info(f"**Age:** {age} years")
            with col3:
                st.info(f"**Header:** {abbreviate_name_age(candidate_name, age)}")
            
            # Convert button
            st.markdown("---")
            
            if st.button("🚀 Convert to Asahi Format", use_container_width=True):
                with st.spinner("🔄 Converting to Asahi format..."):
                    try:
                        final_doc = generate_asahi_cv(raw_text, logo_img, candidate_name, age)
                        
                        # Save to buffer
                        buffer = BytesIO()
                        final_doc.save(buffer)
                        buffer.seek(0)
                        
                        # Success message
                        st.markdown("""
                        <div class="success-message">
                            ✅ CV Successfully Converted to Asahi Format!
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Download button
                        st.download_button(
                            label="⬇️ Download Asahi CV (DOCX)",
                            data=buffer,
                            file_name=f"Asahi_CV_{candidate_name.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                        # Additional info
                        st.markdown("---")
                        st.markdown("""
                        <div style="text-align: center; color: #6b7280; margin-top: 1rem;">
                            <small>📝 Your CV has been formatted according to Asahi standards.</small>
                        </div>
                        """, unsafe_allow_html=True)
                        
                    except Exception as e:
                        st.error(f"❌ Error during conversion: {str(e)}")
                        
        except Exception as e:
            st.error(f"❌ An unexpected error occurred: {str(e)}")
    
    elif uploaded_file or candidate_name or age:
        st.markdown("""
        <div style="background: #fef3c7; border: 1px solid #d97706; border-radius: 8px; padding: 1rem; margin: 1rem 0;">
            <strong>⚠️ Missing Information:</strong><br/>
            Please provide all required information to proceed with CV formatting.
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()